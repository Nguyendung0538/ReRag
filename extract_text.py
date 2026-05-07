import docx
import os
import glob

docs = [
    "chuyen_nhung_co_phan", "dich_vu_sua_chua", "giao_khoan", "hop_tac_kinh_doanh",
    "moi_gioi_mua_ban_bat_dong_san", "nguyen_tac", "phan_phoi", "quang_cao_thuong_mai",
    "tu_van_thiet_ke", "uy_thac_nhap_khau", "uy_thac_xuat_khau", "van_chuyen"
]

base_path = r"d:\Code\ReRag\document\Hop_Dong"

for d in docs:
    p = os.path.join(base_path, d + ".docx")
    try:
        doc = docx.Document(p)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        # find copy count
        copy_idx = text.lower().find("bản")
        
    except Exception as e:
        print(f"Error reading {d}: {e}")

with open("output.txt", "w", encoding="utf-8") as f:
    for d in docs:
        p = os.path.join(base_path, d + ".docx")
        try:
            doc = docx.Document(p)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            f.write(f"--- {d} ---\n")
            if "Tòa án" in text:
                idx = text.find("Tòa án")
                f.write(f"Toa an: {text[max(0, idx-20):idx+50].replace(chr(10), ' ')}\n")
            elif "tòa án" in text.lower():
                idx = text.lower().find("tòa án")
                f.write(f"Toa an: {text[max(0, idx-20):idx+50].replace(chr(10), ' ')}\n")
                
            if "lập thành" in text.lower():
                idx = text.lower().find("lập thành")
                f.write(f"Lap thanh: {text[max(0, idx-10):idx+50].replace(chr(10), ' ')}\n")
                
        except Exception as e:
            f.write(f"Error reading {d}: {e}\n")

